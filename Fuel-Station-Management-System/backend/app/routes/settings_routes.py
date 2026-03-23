from flask import Blueprint, request, jsonify, render_template, send_file, session
from ..models.settings import StationSettings
from .. import db
from datetime import datetime
import os
from werkzeug.utils import secure_filename
from flask import current_app
import io

settings_bp = Blueprint('settings', __name__, url_prefix='/settings')

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_or_create_settings():
    """Always returns the single settings row, creating it if it doesn't exist."""
    s = StationSettings.query.get(1)
    if not s:
        s = StationSettings(
            id             = 1,
            name           = 'BKS Damiru Filling Station',
            address        = 'Pore, Athurugiriya',
            phone          = '',
            email          = '',
            license_number = '',
            logo_path      = '',
        )
        db.session.add(s)
        db.session.commit()
    return s


# ─────────────────────────────────────────────
#  PAGE
# ─────────────────────────────────────────────

@settings_bp.route('/page')
def settings_page():
    return render_template('settings.html')


# ─────────────────────────────────────────────
#  GET STATION SETTINGS
# ─────────────────────────────────────────────

@settings_bp.route('/station', methods=['GET'])
def get_station():
    return jsonify(get_or_create_settings().to_dict())


# ─────────────────────────────────────────────
#  UPDATE STATION SETTINGS
# ─────────────────────────────────────────────

@settings_bp.route('/station', methods=['POST'])
def update_station():
    s    = get_or_create_settings()
    file = request.files.get('logo')

    if file and allowed_file(file.filename):
        filename      = secure_filename(file.filename)
        relative_path = f"uploads/{filename}"
        full_path     = os.path.join(current_app.root_path, 'static', relative_path)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        file.save(full_path)
        s.logo_path = relative_path

    s.name           = request.form.get('name',           s.name)
    s.address        = request.form.get('address',        s.address)
    s.phone          = request.form.get('phone',          s.phone)
    s.email          = request.form.get('email',          s.email)
    s.license_number = request.form.get('license_number', s.license_number)
    s.updated_at     = datetime.utcnow()

    db.session.commit()
    return jsonify({'message': 'Settings saved', 'settings': s.to_dict()})


# ─────────────────────────────────────────────
#  DOWNLOAD LETTERHEAD  (.docx)
# ─────────────────────────────────────────────

@settings_bp.route('/letterhead')
def download_letterhead():
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        return jsonify({'message': 'python-docx not installed. Run: pip install python-docx'}), 500

    s = get_or_create_settings()

    doc = Document()

    # ── Page setup: A4 with narrow margins ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False

    # Clear default header paragraph
    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)

    # Logo + station name side by side using a table in header
    htable = header.add_table(1, 2, Inches(6))
    htable.style = 'Table Grid'

    # Remove borders from header table
    tbl = htable._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Left cell — logo or station initial
    left_cell  = htable.cell(0, 0)
    right_cell = htable.cell(0, 1)
    left_cell.width  = Inches(1.2)
    right_cell.width = Inches(4.8)

    logo_para = left_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo_full_path = None
    if s.logo_path:
        logo_full_path = os.path.join(current_app.root_path, 'static', s.logo_path)

    if logo_full_path and os.path.exists(logo_full_path):
        run = logo_para.add_run()
        run.add_picture(logo_full_path, width=Inches(0.9))
    else:
        # Fallback: station initial in a styled box
        run = logo_para.add_run(s.name[:2].upper() if s.name else 'FS')
        run.font.size  = Pt(20)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x4a)

    # Right cell — station details
    info_para = right_cell.paragraphs[0]
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    name_run = info_para.add_run(s.name or 'Filling Station')
    name_run.font.size  = Pt(14)
    name_run.font.bold  = True
    name_run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x4a)

    details = []
    if s.address:        details.append(s.address)
    if s.phone:          details.append(f"Tel: {s.phone}")
    if s.email:          details.append(f"Email: {s.email}")
    if s.license_number: details.append(f"License: {s.license_number}")

    for detail in details:
        dp = right_cell.add_paragraph()
        dr = dp.add_run(detail)
        dr.font.size  = Pt(9)
        dr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # ── Green divider line under header ──
    divider_para = header.add_paragraph()
    pPr = divider_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a7a4a')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Body ── (large empty area for writing)
    # Add a date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Date: {'_' * 20}")
    date_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Reference line
    ref_para = doc.add_paragraph()
    ref_run = ref_para.add_run("Ref: " + '_' * 30)
    ref_run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Subject line
    subj_para = doc.add_paragraph()
    subj_run = subj_para.add_run("Subject: ")
    subj_run.font.size = Pt(10)
    subj_run.font.bold = True
    subj_line = subj_para.add_run('_' * 50)
    subj_line.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Body placeholder lines
    for _ in range(18):
        lp = doc.add_paragraph()
        lr = lp.add_run('_' * 80)
        lr.font.size  = Pt(10)
        lr.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_run = sig_para.add_run('Authorized Signature')
    sig_run.font.size = Pt(10)
    sig_run.font.bold = True

    sig_line = doc.add_paragraph()
    sig_run2 = sig_line.add_run('_' * 30)
    sig_run2.font.size = Pt(10)

    name_line = doc.add_paragraph()
    name_run2 = name_line.add_run('Name:  ' + '_' * 20)
    name_run2.font.size = Pt(10)

    desig_line = doc.add_paragraph()
    desig_run = desig_line.add_run('Designation:  ' + '_' * 15)
    desig_run.font.size = Pt(10)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)

    # Divider line above footer
    fdivider = footer.add_paragraph()
    fPr = fdivider._p.get_or_add_pPr()
    fBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'),   'single')
    ftop.set(qn('w:sz'),    '6')
    ftop.set(qn('w:space'), '1')
    ftop.set(qn('w:color'), '1a7a4a')
    fBdr.append(ftop)
    fPr.append(fBdr)

    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_parts = []
    if s.name:    footer_parts.append(s.name)
    if s.address: footer_parts.append(s.address)
    if s.phone:   footer_parts.append(f"Tel: {s.phone}")
    footer_run = footer_para.add_run('  |  '.join(footer_parts))
    footer_run.font.size  = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    # ── Save to buffer and send ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    station_slug = (s.name or 'station').replace(' ', '_').lower()
    filename     = f"{station_slug}_letterhead.docx"

    return send_file(
        buffer,
        as_attachment  = True,
        download_name  = filename,
        mimetype       = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )